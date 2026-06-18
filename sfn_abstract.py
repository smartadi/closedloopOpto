from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# ── Title ──────────────────────────────────────────────────────────────────
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title_p.add_run(
    "Controlling cortex in real time: state-dependent limits of "
    "closed-loop optogenetic regulation"
)
r.bold = True
r.font.size = Pt(13)

doc.add_paragraph()

# ── Abstract ───────────────────────────────────────────────────────────────
abstract_text = (
    "Mesoscale cortical activity is typically perturbed open-loop; whether active "
    "feedback control can reliably regulate it—and what limits that control—remains "
    "largely untested. We developed a data-driven, closed-loop optogenetic "
    "controller for population activity and used control-theoretic analysis to show "
    "that its performance is systematically shaped by brain state. In adult "
    "transgenic mice co-expressing GCaMP6 and an inhibitory opsin (n=2 male mice, 13 "
    "sessions; sex differences not assessed; exploratory study), wide-field calcium "
    "imaging (35 Hz) provided real-time ΔF/F feedback over a target cortical region "
    "while a galvanometer-steered laser stimulated an adjacent region. Although "
    "single-trial responses were nonlinear and state-dependent, averaging across "
    "many trials yielded a stable, approximately stationary response well captured "
    "by a low-order linear time-invariant (LTI) model: impulse- and step-response "
    "experiments (0–1.8 mW, ~90 trials/amplitude) fixed the input–output structure "
    "(R²=0.51–0.99 across sessions), from which we derived a proportional-integral "
    "controller with feedforward compensation. Across all 13 sessions, closed-loop "
    "control reduced cross-trial variance by 40.7% during stimulation (signed-rank "
    "p=0.0005; 12/13 sessions) and per-trial tracking error (MSE) by 19.3% "
    "(p=0.0002), while leaving pre-stimulus variance unchanged (p=0.34) — active "
    "disturbance rejection. We then used the controller as a probe of its own "
    "limits: by the internal model principle, a good regulator must implicitly "
    "embody a model of the system, so trials it fails to track mark states where "
    "cortical dynamics depart from that model and become less predictable. Two "
    "independent state variables shaped controllability through opposite mechanisms. "
    "Movement acted as a rejectable disturbance — open-loop tracking degraded as "
    "movement increased whereas closed-loop tracking remained motion-invariant, so "
    "the closed-loop advantage was greatest during movement (interaction p=0.039). "
    "Pre-stimulus ΔF/F variance acted as an irreducible limit: high-variance trials "
    "were tracked worse under both open- and closed-loop control (per-session rank "
    "correlation p=0.0017, 12/13 sessions), because in these less-predictable states "
    "the controller's internal model no longer fits and feedback cannot compensate. "
    "The controller thus rejects what it can model and fails on what it cannot, both "
    "regulating mesoscale cortex and, through its own tracking error, revealing which "
    "brain states are controllable."
)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
r = p.add_run(abstract_text)
r.font.size = Pt(11)
p.paragraph_format.space_after = Pt(0)
p.paragraph_format.line_spacing = Pt(14)

doc.add_paragraph()

# ── Character count ────────────────────────────────────────────────────────
plain = (abstract_text
         .replace('⁻', '-').replace('²', '2').replace('⁶', '6')
         .replace('±', '+-').replace('–', '-').replace('—', '--')
         .replace('Δ', 'D').replace('²', '2').replace('×', 'x'))
chars_no_spaces = len(plain.replace(' ', ''))

note_p = doc.add_paragraph()
nr = note_p.add_run(f"SfN character count (excl. spaces): {chars_no_spaces} / 2,300  |  Headroom: {2300 - chars_no_spaces}")
nr.font.size = Pt(9)
nr.font.color.rgb = RGBColor(0x22, 0x66, 0xBB)
nr.bold = True

doc.add_paragraph()

# ── Verified statistics table ──────────────────────────────────────────────
hdr = doc.add_paragraph()
hdr.add_run("Verified statistics — per-session paired tests (session = unit; n=13, or n=9 for motion)").bold = True
hdr.runs[0].font.size = Pt(10)

table_data = [
    ["Metric", "OL (median)", "CL (median)", "Effect", "Test / p-value"],
    ["Cross-trial variance — stim (0–3 s)",
     "5.27", "2.82",
     "OL/CL = 1.77; −40.7%",
     "Signed-rank  p = 0.0005"],
    ["Cross-trial variance — pre-stim",
     "7.13", "7.45",
     "OL/CL ≈ 1  (no effect)",
     "Signed-rank  p = 0.34"],
    ["Trial MSE, 0–3 s  (session medians)",
     "25.74", "20.77",
     "−19.3%",
     "Signed-rank  p = 0.0002"],
    ["Sessions with OL/CL variance > 1",
     "—", "—", "12 / 13", "—"],
    ["Motion (n=9): OL MSE low→high movement",
     "23.1 → 27.3", "—",
     "OL degrades with movement",
     "per-sess r=+0.25, p=0.055"],
    ["Motion (n=9): CL MSE low→high movement",
     "—", "19.7 → 19.6",
     "CL motion-invariant",
     "per-sess r=+0.11, p=0.36 (n.s.)"],
    ["Motion (n=9): CL advantage larger at high movement",
     "gap +4.1 (low)", "gap +7.0 (high)",
     "interaction (rejectable)",
     "Signed-rank  p = 0.039"],
    ["Pre-stim variance → trial MSE  (OL)",
     "median r = +0.21", "12/13 sess > 0",
     "irreducible limit",
     "Signed-rank  p = 0.0017"],
    ["Pre-stim variance → trial MSE  (CL)",
     "median r = +0.24", "12/13 sess > 0",
     "persists under feedback",
     "Signed-rank  p = 0.0017"],
    ["Motion ⊥ pre-stim variance (independent axes)",
     "r = −0.009", "—",
     "two distinct mechanisms",
     "p = 0.82 (n.s.)"],
    ["[Retired] Delta (1–4 Hz) → trial MSE",
     "—", "—",
     "pooled artifact; within-session null",
     "per-sess p = 0.57"],
]

tbl = doc.add_table(rows=len(table_data), cols=5)
tbl.style = 'Table Grid'
col_widths = [Inches(2.0), Inches(1.1), Inches(1.1), Inches(1.75), Inches(1.8)]

for i, row_data in enumerate(table_data):
    row = tbl.rows[i]
    for j, txt in enumerate(row_data):
        cell = row.cells[j]
        cell.width = col_widths[j]
        run = cell.paragraphs[0].add_run(txt)
        run.font.size = Pt(8)
        if i == 0:
            run.bold = True
            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:color'), 'auto')
            shd.set(qn('w:fill'), 'D6E4F7')
            cell._tc.get_or_add_tcPr().append(shd)

doc.add_paragraph()

# ── Reminders ─────────────────────────────────────────────────────────────
rem = doc.add_paragraph()
rem.add_run("Before submitting  (deadline: June 10, 5 pm EDT):").bold = True
rem.runs[0].font.size = Pt(10)

items = [
    "Add r and p values for 1–4 Hz pre-stim frequency finding — run prestim_variance.m.",
    "Fill author names and department affiliations (main.tex placeholders still open).",
    "Confirm submitting author has active individual SfN membership.",
    "Select nanosymposium preference in the submission form.",
    "Pick theme area: 'Optogenetics: new tools and approaches' or 'Closed-loop / BMI'.",
]
for item in items:
    bp = doc.add_paragraph(style='List Bullet')
    bp.add_run(item).font.size = Pt(10)

import os
output_path = r"C:\Users\aditya\Documents\projects\brain_paper\sfn_abstract.docx"
try:
    doc.save(output_path)
except PermissionError:
    output_path = r"C:\Users\aditya\Documents\projects\brain_paper\sfn_abstract_v2.docx"
    doc.save(output_path)
    print("(primary file was locked/open in Word — saved to sfn_abstract_v2.docx instead)")
print(f"Saved: {output_path}")
print(f"Chars (excl. spaces): {chars_no_spaces}/2300  |  Headroom: {2300-chars_no_spaces}")
